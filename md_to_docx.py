"""
md_to_docx.py — Match Lens Step 5

Converts a markdown report file to a Word document (.docx).
Handles the Match Lens report format: headings, tables, bullet lists, bold text.

Usage:
    from md_to_docx import convert_report, convert_all_reports

    # Convert a single report:
    convert_report("/path/to/tactical_report.md", "/path/to/tactical_report.docx")

    # Convert all reports in a match directory:
    convert_all_reports(MATCH_DIR)

    # Standalone:
    python md_to_docx.py [MATCH_DIR]
    python md_to_docx.py [input.md] [output.docx]
"""

import json
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colours ───────────────────────────────────────────────────────────────────

HEADING1_COLOUR = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
HEADING2_COLOUR = RGBColor(0x2F, 0x5E, 0x8E)   # medium blue
HEADING3_COLOUR = RGBColor(0x40, 0x72, 0xA0)   # lighter blue
TABLE_HEADER_BG = "1F497D"                      # dark blue fill for table headers
TABLE_STRIPE_BG = "EBF3FB"                      # light blue row stripe
BODY_FONT       = "Calibri"
HEADING_FONT    = "Calibri"


# ── Document setup ────────────────────────────────────────────────────────────

def create_document(match_name: str = "") -> Document:
    """Create a new Word document with Match Lens styling."""
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)

    # Override heading styles
    for level, (colour, size) in enumerate([
        (HEADING1_COLOUR, 16),
        (HEADING2_COLOUR, 13),
        (HEADING3_COLOUR, 11.5),
    ], start=1):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name   = HEADING_FONT
        h_style.font.size   = Pt(size)
        h_style.font.color.rgb = colour
        h_style.font.bold   = True

    return doc


# ── Paragraph parsing ─────────────────────────────────────────────────────────

def add_styled_run(para, text: str):
    """
    Add text to a paragraph, applying inline bold/italic as found.
    Handles **bold**, *italic*, `code`, and plain text.
    """
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0

    for match in pattern.finditer(text):
        # Plain text before this match
        plain = text[last_end:match.start()]
        if plain:
            run = para.add_run(plain)
            run.font.name = BODY_FONT

        raw = match.group(0)
        if raw.startswith("**"):
            run = para.add_run(match.group(2))
            run.bold = True
            run.font.name = BODY_FONT
        elif raw.startswith("*"):
            run = para.add_run(match.group(3))
            run.italic = True
            run.font.name = BODY_FONT
        elif raw.startswith("`"):
            run = para.add_run(match.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)

        last_end = match.end()

    # Remaining plain text
    tail = text[last_end:]
    if tail:
        run = para.add_run(tail)
        run.font.name = BODY_FONT


# ── Table parsing ─────────────────────────────────────────────────────────────

def shade_cell(cell, hex_colour: str):
    """Apply background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def add_table(doc: Document, lines: list[str]):
    """
    Parse markdown table lines and add a formatted Word table.
    lines should include header, separator, and data rows.
    """
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-\s|:]+\|\s*$", line):
            continue   # skip separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    n_cols = len(rows[0])
    table  = doc.add_table(rows=len(rows), cols=n_cols)
    table.style             = "Table Grid"
    table.alignment         = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit     = True

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= n_cols:
                break
            cell = row.cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]

            if r_idx == 0:
                # Header row
                shade_cell(cell, TABLE_HEADER_BG)
                run = para.add_run(cell_text)
                run.bold            = True
                run.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name       = HEADING_FONT
                run.font.size       = Pt(9.5)
            else:
                # Data row — stripe even rows
                if r_idx % 2 == 0:
                    shade_cell(cell, TABLE_STRIPE_BG)
                add_styled_run(para, cell_text)
                para.runs[-1].font.size = Pt(9.5) if para.runs else None

    # Space after table
    doc.add_paragraph("")


# ── Horizontal rule ───────────────────────────────────────────────────────────

def add_horizontal_rule(doc: Document):
    """Add a thin horizontal rule (bottom border on an empty paragraph)."""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2F5E8E")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Main conversion ───────────────────────────────────────────────────────────

def convert_report(md_path: str, docx_path: str,
                   match_name: str = "") -> bool:
    """
    Convert a single Markdown report to Word.

    Args:
        md_path:    path to input .md file
        docx_path:  path to output .docx file
        match_name: optional match name for document metadata

    Returns:
        True on success, False on failure
    """
    if not os.path.exists(md_path):
        print(f"  [FAIL] {md_path} not found")
        return False

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    doc = create_document(match_name)
    i   = 0

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # ── Skip empty lines ──────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^[-─=]{3,}\s*$", line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            text  = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # ── Markdown table (starts with |) ────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            add_table(doc, table_lines)
            continue

        # ── Bullet / unordered list ───────────────────────────────────────────
        bullet_match = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if bullet_match:
            indent_len = len(bullet_match.group(1))
            text       = bullet_match.group(2)
            # List depth: 0 = top, 1 = nested
            style_name = "List Bullet" if indent_len == 0 else "List Bullet 2"
            try:
                para = doc.add_paragraph(style=style_name)
            except KeyError:
                para = doc.add_paragraph()
                para.style.paragraph_format.left_indent = Inches(0.25 * (1 + indent_len // 4))
            para.clear()
            add_styled_run(para, text)
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        num_match = re.match(r"^(\s*)\d+[.)]\s+(.*)", line)
        if num_match:
            text = num_match.group(2)
            try:
                para = doc.add_paragraph(style="List Number")
            except KeyError:
                para = doc.add_paragraph()
            para.clear()
            add_styled_run(para, text)
            i += 1
            continue

        # ── Code block ────────────────────────────────────────────────────────
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            if code_lines:
                para = doc.add_paragraph()
                para.style.paragraph_format.left_indent = Inches(0.3)
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            i += 1  # skip closing ```
            continue

        # ── Regular paragraph ─────────────────────────────────────────────────
        para = doc.add_paragraph()
        add_styled_run(para, line.strip())
        i += 1

    # ── Add footer ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    footer  = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    run = footer_para.add_run(
        f"Match Lens  |  {match_name}  |  "
        f"Generated {datetime.now().strftime('%d %b %Y')}"
    )
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        doc.save(docx_path)
        print(f"  [OK] {os.path.basename(md_path)} → {os.path.basename(docx_path)}")
        return True
    except Exception as e:
        print(f"  [FAIL] Failed to save {docx_path}: {e}")
        return False


# ── Batch conversion ──────────────────────────────────────────────────────────

REPORT_FILES = [
    "tactical_report.md",
    "opposition_report_*.md",
]


def convert_all_reports(match_dir: str) -> dict:
    """
    Convert all markdown report files in match_dir to Word documents.
    Returns summary of conversions.
    """
    import glob as globmod

    # Load match name for footer
    match_name = ""
    config_path = os.path.join(match_dir, "match_config.json")
    if os.path.exists(config_path):
        with open(config_path, encoding="utf-8") as f:
            config = json.load(f)
        match_name = config.get("match", "")

    print(f"\n{'-'*55}")
    print(f"  Step 5 — Converting reports to Word")
    print(f"  Match: {match_name or match_dir}")
    print(f"{'-'*55}")

    converted = 0
    failed    = 0
    skipped   = 0

    # Find all markdown reports
    md_files = []
    for pattern in REPORT_FILES:
        if "*" in pattern:
            md_files.extend(globmod.glob(os.path.join(match_dir, pattern)))
        else:
            candidate = os.path.join(match_dir, pattern)
            if os.path.exists(candidate):
                md_files.append(candidate)

    if not md_files:
        print(f"  No report .md files found in {match_dir}")
        return {"converted": 0, "failed": 0, "skipped": 0}

    for md_path in sorted(md_files):
        docx_path = md_path.replace(".md", ".docx")
        if os.path.exists(docx_path):
            md_mtime   = os.path.getmtime(md_path)
            docx_mtime = os.path.getmtime(docx_path)
            if docx_mtime >= md_mtime:
                print(f"  → {os.path.basename(docx_path)} up to date — skipping")
                skipped += 1
                continue

        success = convert_report(md_path, docx_path, match_name)
        if success:
            converted += 1
        else:
            failed += 1

    print(f"{'-'*55}")
    print(f"  Converted: {converted}  Failed: {failed}  Skipped: {skipped}")
    print(f"{'-'*55}")

    return {"converted": converted, "failed": failed, "skipped": skipped}


if __name__ == "__main__":
    if len(sys.argv) == 3:
        # Direct conversion: python md_to_docx.py input.md output.docx
        convert_report(sys.argv[1], sys.argv[2])
    elif len(sys.argv) == 2:
        # Batch: python md_to_docx.py MATCH_DIR
        convert_all_reports(sys.argv[1])
    else:
        match_dir = input("Match directory: ").strip()
        convert_all_reports(match_dir)
